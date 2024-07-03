import docx

# Create a new Document
doc = docx.Document()

# Add a title
doc.add_heading('CAN Packet Data Structure', 0)

# Add table with appropriate headings and subheadings
table = doc.add_table(rows=1, cols=14)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'CAN Packet'
hdr_cells[1].text = 'Frame ID'
hdr_cells[2].text = 'Extended Frame'
hdr_cells[3].text = 'Length'
hdr_cells[4].text = 'Signal'
hdr_cells[5].text = 'Start Bit'
hdr_cells[6].text = 'Length'
hdr_cells[7].text = 'Byte Order'
hdr_cells[8].text = 'Signed'
hdr_cells[9].text = 'Scale'
hdr_cells[10].text = 'Offset'
hdr_cells[11].text = 'Minimum'
hdr_cells[12].text = 'Maximum'
hdr_cells[13].text = 'Unit'

# Define the data structure
data = [
    ("CAN_PACKET_BMS_TEMPS", 11018, "Yes", 8, "NoOfCells", 15, 8, "Big Endian", "No", 1, 0, 0, 1, ""),
    ("", "", "", "", "auxVoltagesIndividual1", 23, 16, "Big Endian", "No", 0.01, 0, "", "", ""),
    ("", "", "", "", "auxVoltagesIndividual2", 39, 16, "Big Endian", "No", 0.01, 0, "", "", ""),
    ("", "", "", "", "auxVoltagesIndividual3", 55, 16, "Big Endian", "No", 0.01, 0, "", "", ""),
    ("CAN_PACKET_BMS_V_TOT", 9738, "Yes", 8, "packVoltage", 7, 32, "Big Endian", "Yes", 1, 0, -3.4e38, 3.4e38, "V"),
    ("", "", "", "", "chargerVoltage", 39, 32, "Big Endian", "Yes", 1, 0, 0, 1, "V"),
    ("CAN_PACKET_BMS_I", 9994, "Yes", 8, "packCurrent1", 7, 32, "Big Endian", "Yes", 1, 0, 0, 1, "A"),
    ("", "", "", "", "packCurrent2", 39, 32, "Big Endian", "Yes", 1, 0, 0, 1, "A"),
    ("CAN_PACKET_BMS_AH_WH", 10250, "Yes", 8, "Ah_Counter", 7, 32, "Big Endian", "Yes", 1, 0, 0, 1, "Ah"),
    ("", "", "", "", "Wh_Counter", 39, 32, "Big Endian", "Yes", 1, 0, 0, 1, "Wh"),
    ("CAN_PACKET_BMS_V_CELL", 10506, "Yes", 8, "cellPoint", 7, 8, "Big Endian", "No", 1, 0, 0, 48, ""),
    ("", "", "", "", "NoOfCells", 15, 8, "Big Endian", "No", 1, 0, 0, 48, ""),
    ("", "", "", "", "cellVoltage10", 23, 16, "Big Endian", "No", 0.001, 0, 0, 1, "V"),
    ("", "", "", "", "cellVoltage11", 39, 16, "Big Endian", "No", 0.001, 0, 0, 1, "V"),
    ("", "", "", "", "cellVoltage12", 55, 16, "Big Endian", "No", 0.001, 0, 0, 1, "V"),
    ("CAN_PACKET_BMS_BAL", 10762, "Yes", 8, "NoOfCells", 0, 8, "Little Endian", "Yes", 1, 0, "", "", ""),
    ("", "", "", "", "bal_state", 8, 56, "Little Endian", "Yes", 1, 0, 0, 1, ""),
    ("CAN_PACKET_BMS_SOC_SOH_TEMP_STAT", 11530, "Yes", 8, "cellVoltageLow", 7, 16, "Big Endian", "No", 0.001, 0, 0, 1, "V"),
    ("", "", "", "", "cellVoltageHigh", 23, 16, "Big Endian", "No", 0.001, 0, 0, 1, "V"),
    ("", "", "", "", "SOC", 39, 8, "Big Endian", "No", 0.392156863, 0, 0, 1, "%"),
    ("", "", "", "", "SOH", 47, 8, "Big Endian", "No", 0.3922, 0, 0, 100, "%"),
    ("", "", "", "", "tBattHi", 55, 8, "Big Endian", "No", 1, 0, 0, 1, "C"),
    ("", "", "", "", "BitF", 63, 8, "Big Endian", "No", 1, 0, 0, 1, ""),
    ("CAN_PACKET_BMS_HUM", 11274, "Yes", 6, "CAN_PACKET_BMS_TEMP0", 7, 16, "Big Endian", "No", 0.01, 0, 0, 1, "C"),
    ("", "", "", "", "CAN_PACKET_BMS_HUM_HUM", 23, 16, "Big Endian", "No", 0.01, 0, 0, 1, "%"),
    ("", "", "", "", "CAN_PACKET_BMS_HUM_TEMP1", 39, 16, "Big Endian", "No", 0.01, 0, 0, 1, "C"),
]

# Add data to the table
for row in data:
    cells = table.add_row().cells
    for i, item in enumerate(row):
        cells[i].text = str(item)

# Save the document
file_path = 'CAN_Packet_Data_Structure.docx'
doc.save(file_path)
