from docx import Document

# Create a new Document
doc = Document()

# Add a title
doc.add_heading('AT Command Reference', 0)

# Create a table
table = doc.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Command'
hdr_cells[1].text = 'Description'

# List of commands and descriptions
commands = [
    ("+++\\r\\n", "Enter AT command mode"),
    ("AT+EXIT\\r\\n", "Exit AT command mode"),
    ("ATE\r\n", "Enable/disable AT command echo"),
    ("AT+VER\r\n", "Check the software version number"),
    ("AT+HELP\r\n", "View AT help"),
    ("AT+SF=7\r\n", "Set the spreading factor to 7, the value range is 7~12"),
    ("AT+BW=0\r\n", "Set bandwidth, 0 means 125KHz, 1 means 250KHz, 2 means 500KHz"),
    ("AT+CR=1\r\n", "Set the encoding rate to 1, 1 represents 4/5, 2 represents 4/6, 3 represents 4/7, 4 represents 4/8"),
    ("AT+PWR=22\r\n", "Set the RF power, the value range is 10~22dBm"),
    ("AT+NETID=0\r\n", "Network ID assignment, the value range is 0~65535"),
    ("AT+LBT=0\r\n", "Enable/disable LBT function, 0: disable, 1: enable"),
    ("AT+MODE=1", "DTU working mode, 1: stream mode, 2: packet mode, 3: relay mode"),
    ("AT+TXCH=18\r\n", "Transmit channel, value range 0~80, corresponding frequency point is 850~930MHz or 410~490MHz"),
    ("AT+RXCH=18\r\n", "Receive channel, value range 0~80, corresponding frequency point is 850~930MHz or 410~490MHz"),
    ("AT+RSSI=0\r\n", "Enable/disable RSSI signal value output, 0: disable, 1: enable"),
    ("AT+ADDR=0\r\n", "Set DTU address, value range 0~65535"),
    ("AT+PORT=3\r\n", "Set COM port, 1:RS422, 2:RS485, 3:RS232"),
    ("AT+BAUD=115200\r\n", "Set COMx port baud rate, value range 1200~115200, 1200, 2400, ....., 57600, 115200"),
    ("AT+COMM=\"8N1\"\r\n", "Set COM port parameters, data bits: 8 or 9, parity: N, O, E, stop bits: 0, 1, 2"),
    ("AT+AllP=7,125,1,22,0,0,1,18,18,0,0,3,115200, \"8N1\",0", "Set the spreading factor to key multi-parameter"),
    ("AT+RESTORE=0\r\n", "Restore factory settings, 0: disabled, 1: enabled")
]

# Populate the table with commands and descriptions
for cmd, desc in commands:
    row_cells = table.add_row().cells
    row_cells[0].text = cmd
    row_cells[1].text = desc

# Save the document
doc.save("AT_Command_Reference.docx")